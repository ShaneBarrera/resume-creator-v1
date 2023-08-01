from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

name = "Silas Raye"
email = "silas.raye@gmail.com"
phone = "352-231-0747"
edu = "Bachelor's degree in data science from the University of Florida"

# Create a new Document
doc = Document()

# Set the margins
section = doc.sections[0]
section.left_margin = Inches(0.25)
section.right_margin = Inches(0.25)
section.top_margin = Inches(0.25)
section.bottom_margin = Inches(0.25)

# Add a title
title = doc.add_heading(level=1)
run = title.add_run(name.upper())
run.bold = True
run.font.size = Pt(24)  # Increase font size
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add a subtitle
subtitle = doc.add_paragraph(email + " | " + phone)
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add an underlined heading
education_heading = doc.add_heading('EDUCATION', level=2)
run = education_heading.runs[0]
run.underline = True

doc.add_paragraph(edu, style='Body Text')
# doc.add_paragraph('Bullet 2', style='List Bullet')
# doc.add_paragraph('Bullet 3', style='List Bullet')


# Add an underlined heading
skills_heading = doc.add_heading('TECHNICAL SKILLS', level=2)
run = skills_heading.runs[0]
run.underline = True

# Add an underlined heading
exp_heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
run = exp_heading.runs[0]
run.underline = True

# Save the document
doc.save('dummy.docx')

print('done')
